from __future__ import annotations

from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Any


@dataclass
class ParsedNode:
    type: str
    text: str
    page: int | None = None
    level: int | None = None
    section_path: list[str] | None = None


@dataclass
class ParsedDocument:
    file_type: str
    nodes: list[ParsedNode]
    page_count: int = 0
    paragraph_count: int = 0
    table_count: int = 0
    image_count: int = 0

    def to_dict(self) -> dict[str, Any]:
        return {
            "fileType": self.file_type,
            "pageCount": self.page_count,
            "paragraphCount": self.paragraph_count,
            "tableCount": self.table_count,
            "imageCount": self.image_count,
            "nodes": [asdict(node) for node in self.nodes],
        }


def parse_document(file_path: str, file_type: str | None = None) -> ParsedDocument:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {path}")
    detected_type = (file_type or path.suffix.lstrip(".")).upper()
    if detected_type == "PDF":
        return parse_pdf(path)
    if detected_type == "DOCX":
        return parse_docx(path)
    if detected_type in ("MD", "MARKDOWN", "TXT"):
        return parse_text(path, detected_type)
    raise ValueError(f"不支持的文档类型: {detected_type}")


def parse_pdf(path: Path) -> ParsedDocument:
    try:
        import fitz  # type: ignore
    except Exception as exc:
        raise RuntimeError("PDF 解析依赖 PyMuPDF 未安装，请安装 pymupdf") from exc

    nodes: list[ParsedNode] = []
    doc = fitz.open(path)
    try:
        for page_index in range(doc.page_count):
            page = doc.load_page(page_index)
            text = page.get_text("text").strip()
            if not text:
                continue
            for paragraph in split_paragraphs(text):
                nodes.append(ParsedNode(type="paragraph", text=paragraph, page=page_index + 1))
        return ParsedDocument(
            file_type="PDF",
            nodes=nodes,
            page_count=doc.page_count,
            paragraph_count=len(nodes),
        )
    finally:
        doc.close()


def parse_docx(path: Path) -> ParsedDocument:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise RuntimeError("DOCX 解析依赖 python-docx 未安装，请安装 python-docx") from exc

    document = Document(path)
    nodes: list[ParsedNode] = []
    section_path: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name.startswith("Heading"):
            level = heading_level(style_name)
            section_path = section_path[: max(0, level - 1)]
            section_path.append(text)
            nodes.append(ParsedNode(type="heading", text=text, level=level, section_path=list(section_path)))
        else:
            nodes.append(ParsedNode(type="paragraph", text=text, section_path=list(section_path)))

    table_count = 0
    for table in document.tables:
        table_count += 1
        markdown = table_to_markdown(table)
        if markdown:
            nodes.append(ParsedNode(type="table", text=markdown, section_path=list(section_path)))

    paragraph_count = len([node for node in nodes if node.type in ("heading", "paragraph")])
    return ParsedDocument(
        file_type="DOCX",
        nodes=nodes,
        paragraph_count=paragraph_count,
        table_count=table_count,
    )


def parse_text(path: Path, file_type: str) -> ParsedDocument:
    text = path.read_text(encoding="utf-8")
    nodes = [ParsedNode(type="paragraph", text=paragraph) for paragraph in split_paragraphs(text)]
    return ParsedDocument(file_type=file_type, nodes=nodes, paragraph_count=len(nodes))


def split_paragraphs(text: str) -> list[str]:
    return [
        paragraph.strip()
        for paragraph in text.replace("\r\n", "\n").split("\n")
        if paragraph.strip()
    ]


def heading_level(style_name: str) -> int:
    tail = style_name.replace("Heading", "").strip()
    try:
        return max(1, min(int(tail), 6))
    except ValueError:
        return 1


def table_to_markdown(table: Any) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([cell.text.strip().replace("\n", " ") for cell in row.cells])
    if not rows:
        return ""
    header = rows[0]
    separator = ["---" for _ in header]
    body = rows[1:]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    lines.extend("| " + " | ".join(row) + " |" for row in body)
    return "\n".join(lines)
